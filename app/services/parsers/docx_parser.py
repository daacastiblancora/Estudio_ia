from typing import List
from langchain_core.documents import Document
from app.services.parser_registry import BaseParser
from app.core.logging import logger

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


class DOCXParser(BaseParser):
    """DOCX parser using python-docx — extracts paragraphs with section tracking."""
    supported_extensions = ["docx"]

    def parse(self, file_path: str, source_name: str) -> List[Document]:
        if DocxDocument is None:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return []

        doc = DocxDocument(file_path)
        documents = []
        current_section = "General"
        section_num = 1
        section_texts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect section headings (Heading styles)
            if para.style and para.style.name and "Heading" in para.style.name:
                # Flush previous section
                if section_texts:
                    content = self._inject_header(source_name, "Sección", section_num)
                    content += f"Sección: {current_section}\n\n"
                    content += "\n".join(section_texts)

                    documents.append(Document(
                        page_content=content,
                        metadata=self._make_metadata(
                            source=source_name, page=section_num,
                            total=0, fmt="docx", section=current_section
                        ),
                    ))
                    section_num += 1
                    section_texts = []

                current_section = text
            else:
                section_texts.append(text)

        # Flush last section
        if section_texts:
            content = self._inject_header(source_name, "Sección", section_num)
            content += f"Sección: {current_section}\n\n"
            content += "\n".join(section_texts)

            documents.append(Document(
                page_content=content,
                metadata=self._make_metadata(
                    source=source_name, page=section_num,
                    total=section_num, fmt="docx", section=current_section
                ),
            ))

        # Update total_pages in all docs
        for d in documents:
            d.metadata["total_pages"] = section_num

        return documents
